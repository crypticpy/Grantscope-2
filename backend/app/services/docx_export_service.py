"""DOCX export service for grant proposals, budgets, and application packages.

Generates professional Word documents using python-docx with City of Austin
branding, proper formatting, and structured content from proposals, budget
line items, and checklist data.
"""

import io
import logging
from datetime import datetime, timezone
from decimal import Decimal
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

_STANDARD_SECTIONS = [
    "executive_summary",
    "needs_statement",
    "project_description",
    "budget_narrative",
    "timeline",
    "evaluation_plan",
]

_FEDERAL_BUDGET_CATEGORIES = [
    "personnel",
    "fringe_benefits",
    "travel",
    "equipment",
    "supplies",
    "contractual",
    "construction",
    "other",
    "indirect_costs",
]

_HEADER_BG = RGBColor(0x4A, 0x5A, 0x6A)  # Slate gray for table headers
_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)  # White text on header rows
_SUBTOTAL_BG = RGBColor(0xE8, 0xEC, 0xF0)  # Light gray for subtotal rows
_BRAND_COLOR = RGBColor(0x1A, 0x3C, 0x5E)  # City of Austin navy


class DocxExportService:
    """Static methods for generating DOCX exports of grant application data."""

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _set_margins(document: Document, inches: float = 1.0) -> None:
        """Set all page margins to the specified value in inches."""
        for section in document.sections:
            section.top_margin = Inches(inches)
            section.bottom_margin = Inches(inches)
            section.left_margin = Inches(inches)
            section.right_margin = Inches(inches)

    @staticmethod
    def _set_default_font(
        document: Document,
        name: str = "Times New Roman",
        size: int = 12,
    ) -> None:
        """Set the default paragraph font for the entire document."""
        style = document.styles["Normal"]
        font = style.font
        font.name = name
        font.size = Pt(size)
        font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

        # Also update heading styles to use a complementary font
        for level in range(1, 4):
            heading_name = f"Heading {level}"
            if heading_name in document.styles:
                hstyle = document.styles[heading_name]
                hstyle.font.name = name
                hstyle.font.color.rgb = _BRAND_COLOR

    @staticmethod
    def _add_table_with_header(
        document: Document,
        headers: list[str],
        widths: list[float] | None = None,
    ):
        """Create a styled table with a bold, gray-background header row.

        Args:
            document: The Document to add the table to.
            headers: Column header strings.
            widths: Optional column widths in inches.

        Returns:
            The created Table object.
        """
        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Style the header row
        header_row = table.rows[0]
        for idx, header_text in enumerate(headers):
            cell = header_row.cells[idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(header_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = _HEADER_TEXT

            # Set cell background color
            shading = cell._element.get_or_add_tcPr()
            shading_elem = shading.makeelement(
                qn("w:shd"),
                {
                    qn("w:fill"): "4A5A6A",
                    qn("w:val"): "clear",
                },
            )
            shading.append(shading_elem)

        # Apply column widths if provided
        if widths:
            for row in table.rows:
                for idx, width in enumerate(widths):
                    if idx < len(row.cells):
                        row.cells[idx].width = Inches(width)

        return table

    @staticmethod
    def _format_currency(amount: Any) -> str:
        """Format a numeric value as USD currency string.

        Args:
            amount: A number, Decimal, string, or None.

        Returns:
            Formatted string like "$1,234.56" or "$0.00" for invalid input.
        """
        if amount is None:
            return "$0.00"
        try:
            value = float(amount)
            return f"${value:,.2f}"
        except (ValueError, TypeError):
            return "$0.00"

    @staticmethod
    def _add_page_headers_and_footers(
        document: Document,
        department: str = "APH",
    ) -> None:
        """Add page headers and footers to every section of the document.

        Header: "City of Austin -- [Department]" right-aligned.
        Footer: Page number centered, "Prepared with GrantScope" right-aligned.
        """
        for section in document.sections:
            # Header
            header = section.header
            header.is_linked_to_previous = False
            header_para = (
                header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            )
            header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = header_para.add_run(f"City of Austin \u2014 {department}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            run.italic = True

            # Footer
            footer = section.footer
            footer.is_linked_to_previous = False

            # Clear existing paragraphs
            for p in footer.paragraphs:
                p.clear()

            # Page number centered
            page_para = (
                footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            )
            page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = page_para.add_run("Page ")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            # Insert PAGE field
            fld_char_begin = run._element.makeelement(
                qn("w:fldChar"), {qn("w:fldCharType"): "begin"}
            )
            run._element.addnext(fld_char_begin)
            instr_text = run._element.makeelement(qn("w:instrText"), {})
            instr_text.text = " PAGE "
            fld_char_begin.addnext(instr_text)
            fld_char_end = run._element.makeelement(
                qn("w:fldChar"), {qn("w:fldCharType"): "end"}
            )
            instr_text.addnext(fld_char_end)

            # "Prepared with GrantScope" right-aligned
            brand_para = footer.add_paragraph()
            brand_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            brand_run = brand_para.add_run("Prepared with GrantScope")
            brand_run.font.size = Pt(8)
            brand_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            brand_run.italic = True

    @staticmethod
    def _section_title(key: str) -> str:
        """Convert a section key like 'executive_summary' to 'Executive Summary'."""
        return key.replace("_", " ").title()

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    @staticmethod
    def generate_proposal_docx(
        proposal: dict,
        grant_context: dict | None = None,
        plan_data: dict | None = None,
        budget_items: list | None = None,
    ) -> bytes:
        """Generate a complete proposal DOCX document.

        Args:
            proposal: Proposal dict with title, sections, etc.
            grant_context: Optional grant context with grantor name, deadline, etc.
            plan_data: Optional plan data which may contain checklist_items.
            budget_items: Optional list of budget line item dicts.

        Returns:
            The DOCX file contents as bytes.
        """
        doc = Document()
        svc = DocxExportService

        svc._set_margins(doc, inches=1.0)
        svc._set_default_font(doc, name="Times New Roman", size=12)

        department = "City of Austin"
        if grant_context and grant_context.get("department"):
            department = grant_context["department"]

        svc._add_page_headers_and_footers(doc, department=department)

        # ---- Cover Page ----
        # Add some spacing before the title
        for _ in range(4):
            doc.add_paragraph("")

        grant_name = proposal.get("title", "Grant Proposal")
        if grant_context and grant_context.get("grant_name"):
            grant_name = grant_context["grant_name"]

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(grant_name)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = _BRAND_COLOR

        doc.add_paragraph("")

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("Applicant: City of Austin")
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        if department and department != "City of Austin":
            dept_para = doc.add_paragraph()
            dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            dept_run = dept_para.add_run(f"Department: {department}")
            dept_run.font.size = Pt(12)
            dept_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Horizontal rule
        doc.add_paragraph("")
        hr_para = doc.add_paragraph()
        hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr_run = hr_para.add_run("_" * 60)
        hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        # Page break after cover
        doc.add_page_break()

        # ---- Table of Contents Placeholder ----
        doc.add_heading("TABLE OF CONTENTS", level=1)
        toc_para = doc.add_paragraph()
        toc_run = toc_para.add_run("[Update field to generate]")
        toc_run.italic = True
        toc_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        doc.add_page_break()

        # ---- Standard Sections ----
        sections = proposal.get("sections") or {}

        for section_key in _STANDARD_SECTIONS:
            section_data = sections.get(section_key)
            if not section_data:
                continue

            # Extract content string
            if isinstance(section_data, dict):
                content = section_data.get("content", "")
            elif isinstance(section_data, str):
                content = section_data
            else:
                continue

            if not content or not content.strip():
                continue

            doc.add_heading(svc._section_title(section_key), level=1)

            # Split on double newlines for paragraph breaks
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                stripped = para_text.strip()
                if stripped:
                    doc.add_paragraph(stripped)

        # Also include any custom/extra sections not in the standard list
        for section_key, section_data in sections.items():
            if section_key in _STANDARD_SECTIONS:
                continue

            if isinstance(section_data, dict):
                content = section_data.get("content", "")
            elif isinstance(section_data, str):
                content = section_data
            else:
                continue

            if not content or not content.strip():
                continue

            doc.add_heading(svc._section_title(section_key), level=1)
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                stripped = para_text.strip()
                if stripped:
                    doc.add_paragraph(stripped)

        # ---- Budget Summary Table ----
        if budget_items:
            doc.add_page_break()
            doc.add_heading("Budget Summary", level=1)

            headers = ["Category", "Description", "Total", "Federal", "Match"]
            widths = [1.5, 2.5, 1.0, 1.0, 1.0]
            table = svc._add_table_with_header(doc, headers, widths)

            # Group items by category for subtotals
            category_totals: dict[str, float] = {}
            grand_total = 0.0
            grand_federal = 0.0
            grand_match = 0.0

            for item in budget_items:
                cat = item.get("category", "other")
                total = float(item.get("total_cost", 0) or 0)
                federal = float(item.get("federal_share", 0) or 0)
                match = float(item.get("match_share", 0) or 0)

                category_totals[cat] = category_totals.get(cat, 0) + total
                grand_total += total
                grand_federal += federal
                grand_match += match

                row = table.add_row()
                row.cells[0].text = svc._section_title(cat)
                row.cells[1].text = item.get("description", "")
                row.cells[2].text = svc._format_currency(total)
                row.cells[3].text = svc._format_currency(federal)
                row.cells[4].text = svc._format_currency(match)

            # Category subtotal rows
            for cat, cat_total in category_totals.items():
                if len([i for i in budget_items if i.get("category") == cat]) > 1:
                    row = table.add_row()
                    row.cells[0].text = f"{svc._section_title(cat)} Subtotal"
                    row.cells[1].text = ""
                    row.cells[2].text = svc._format_currency(cat_total)
                    row.cells[3].text = ""
                    row.cells[4].text = ""
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.italic = True

            # Grand total row (bold)
            total_row = table.add_row()
            total_row.cells[0].text = "GRAND TOTAL"
            total_row.cells[1].text = ""
            total_row.cells[2].text = svc._format_currency(grand_total)
            total_row.cells[3].text = svc._format_currency(grand_federal)
            total_row.cells[4].text = svc._format_currency(grand_match)
            for cell in total_row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # ---- Appendix: Requirements Checklist ----
        checklist_items = None
        if plan_data and plan_data.get("checklist_items"):
            checklist_items = plan_data["checklist_items"]

        if checklist_items:
            doc.add_page_break()
            doc.add_heading("Appendix: Requirements Checklist", level=1)

            cl_headers = ["Description", "Status", "Mandatory"]
            cl_widths = [4.0, 1.5, 1.0]
            cl_table = svc._add_table_with_header(doc, cl_headers, cl_widths)

            for item in checklist_items:
                row = cl_table.add_row()
                row.cells[0].text = item.get("description", "")
                is_completed = item.get("is_completed", False)
                row.cells[1].text = "Complete" if is_completed else "Incomplete"
                is_mandatory = item.get("is_mandatory", False)
                row.cells[2].text = "Yes" if is_mandatory else "No"

        # ---- Save to bytes ----
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def generate_budget_docx(
        budget_items: list,
        settings: dict,
        calculations: dict,
    ) -> bytes:
        """Generate a standalone budget detail DOCX document.

        Args:
            budget_items: List of budget line item dicts.
            settings: Budget settings dict (fringe_rate, indirect_rate, etc.).
            calculations: Pre-computed budget calculations (direct_total,
                indirect_total, grand_total, federal_share, match_share).

        Returns:
            The DOCX file contents as bytes.
        """
        doc = Document()
        svc = DocxExportService

        svc._set_margins(doc, inches=1.0)
        svc._set_default_font(doc, name="Times New Roman", size=12)
        svc._add_page_headers_and_footers(doc, department="City of Austin")

        # Title
        title = calculations.get("title", "Grant Application")
        title_para = doc.add_heading(f"Budget Detail \u2014 {title}", level=1)

        doc.add_paragraph(
            f"Prepared: {datetime.now(timezone.utc).strftime('%B %d, %Y')}"
        )
        doc.add_paragraph("")

        # Group items by category
        items_by_category: dict[str, list[dict]] = {}
        for item in budget_items:
            cat = item.get("category", "other")
            items_by_category.setdefault(cat, []).append(item)

        # Section per category
        for category in _FEDERAL_BUDGET_CATEGORIES:
            cat_items = items_by_category.get(category, [])
            if not cat_items:
                continue

            doc.add_heading(svc._section_title(category), level=2)

            if category == "personnel":
                # Personnel table has special columns
                headers = ["Role", "FTE", "Salary", "Months", "Total"]
                widths = [2.0, 0.8, 1.2, 0.8, 1.2]
                table = svc._add_table_with_header(doc, headers, widths)

                cat_total = 0.0
                for item in cat_items:
                    row = table.add_row()
                    row.cells[0].text = item.get("role", item.get("description", ""))
                    row.cells[1].text = str(item.get("fte", ""))
                    row.cells[2].text = svc._format_currency(item.get("annual_salary"))
                    row.cells[3].text = str(item.get("months_on_project", ""))
                    row.cells[4].text = svc._format_currency(item.get("total_cost"))
                    cat_total += float(item.get("total_cost", 0) or 0)

                # Subtotal row
                sub_row = table.add_row()
                sub_row.cells[0].text = "Personnel Subtotal"
                sub_row.cells[4].text = svc._format_currency(cat_total)
                for cell in sub_row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
            else:
                # Generic table
                headers = ["Description", "Qty", "Unit Cost", "Total", "Justification"]
                widths = [2.0, 0.6, 1.0, 1.0, 2.0]
                table = svc._add_table_with_header(doc, headers, widths)

                cat_total = 0.0
                for item in cat_items:
                    row = table.add_row()
                    row.cells[0].text = item.get("description", "")
                    row.cells[1].text = str(item.get("quantity", ""))
                    row.cells[2].text = svc._format_currency(item.get("unit_cost"))
                    row.cells[3].text = svc._format_currency(item.get("total_cost"))
                    row.cells[4].text = item.get("justification", "") or ""
                    cat_total += float(item.get("total_cost", 0) or 0)

                # Subtotal row
                sub_row = table.add_row()
                sub_row.cells[0].text = f"{svc._section_title(category)} Subtotal"
                sub_row.cells[3].text = svc._format_currency(cat_total)
                for cell in sub_row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

            doc.add_paragraph("")  # spacing between categories

        # Also handle any categories not in the standard list
        for category, cat_items in items_by_category.items():
            if category in _FEDERAL_BUDGET_CATEGORIES:
                continue

            doc.add_heading(svc._section_title(category), level=2)
            headers = ["Description", "Qty", "Unit Cost", "Total", "Justification"]
            widths = [2.0, 0.6, 1.0, 1.0, 2.0]
            table = svc._add_table_with_header(doc, headers, widths)

            cat_total = 0.0
            for item in cat_items:
                row = table.add_row()
                row.cells[0].text = item.get("description", "")
                row.cells[1].text = str(item.get("quantity", ""))
                row.cells[2].text = svc._format_currency(item.get("unit_cost"))
                row.cells[3].text = svc._format_currency(item.get("total_cost"))
                row.cells[4].text = item.get("justification", "") or ""
                cat_total += float(item.get("total_cost", 0) or 0)

            sub_row = table.add_row()
            sub_row.cells[0].text = f"{svc._section_title(category)} Subtotal"
            sub_row.cells[3].text = svc._format_currency(cat_total)
            for cell in sub_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            doc.add_paragraph("")

        # ---- Summary Section ----
        doc.add_page_break()
        doc.add_heading("Budget Summary", level=1)

        direct_total = calculations.get("direct_total", 0)
        indirect_total = calculations.get("indirect_total", 0)
        grand_total = calculations.get("grand_total", 0)

        summary_headers = ["Item", "Amount"]
        summary_widths = [4.0, 2.0]
        summary_table = svc._add_table_with_header(doc, summary_headers, summary_widths)

        # Direct costs
        row = summary_table.add_row()
        row.cells[0].text = "Direct Costs Subtotal"
        row.cells[1].text = svc._format_currency(direct_total)

        # Indirect costs
        indirect_rate = settings.get("indirect_rate", 0)
        indirect_base = settings.get("indirect_base", "mtdc")
        row = summary_table.add_row()
        rate_display = (
            f"{float(indirect_rate or 0) * 100:.1f}%" if indirect_rate else "N/A"
        )
        row.cells[0].text = (
            f"Indirect Costs ({rate_display} on {indirect_base.upper()})"
        )
        row.cells[1].text = svc._format_currency(indirect_total)

        # Grand total
        total_row = summary_table.add_row()
        total_row.cells[0].text = "GRAND TOTAL"
        total_row.cells[1].text = svc._format_currency(grand_total)
        for cell in total_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        # Match breakdown if applicable
        match_required = settings.get("match_required", False)
        if match_required:
            doc.add_paragraph("")
            doc.add_heading("Cost Sharing / Match", level=2)

            match_headers = ["Source", "Amount"]
            match_widths = [4.0, 2.0]
            match_table = svc._add_table_with_header(doc, match_headers, match_widths)

            federal_share = calculations.get("federal_share", grand_total)
            match_share = calculations.get("match_share", 0)

            row = match_table.add_row()
            row.cells[0].text = "Federal Share"
            row.cells[1].text = svc._format_currency(federal_share)

            row = match_table.add_row()
            row.cells[0].text = "Applicant Match"
            row.cells[1].text = svc._format_currency(match_share)

            total_row = match_table.add_row()
            total_row.cells[0].text = "Total Project Cost"
            total_row.cells[1].text = svc._format_currency(
                float(federal_share or 0) + float(match_share or 0)
            )
            for cell in total_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

        # ---- Save to bytes ----
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def generate_package_docx(
        proposal: dict,
        budget_items: list,
        budget_settings: dict,
        budget_calculations: dict,
        checklist_items: list,
        grant_context: dict | None = None,
    ) -> bytes:
        """Generate a combined application package DOCX (proposal + budget + checklist).

        Produces a single document with page breaks between each major section:
        the full proposal narrative, the detailed budget, and the requirements
        checklist.

        Args:
            proposal: Proposal dict with title and sections.
            budget_items: List of budget line item dicts.
            budget_settings: Budget settings dict.
            budget_calculations: Pre-computed budget calculations.
            checklist_items: List of checklist item dicts.
            grant_context: Optional grant context metadata.

        Returns:
            The DOCX file contents as bytes.
        """
        doc = Document()
        svc = DocxExportService

        svc._set_margins(doc, inches=1.0)
        svc._set_default_font(doc, name="Times New Roman", size=12)

        department = "City of Austin"
        if grant_context and grant_context.get("department"):
            department = grant_context["department"]

        svc._add_page_headers_and_footers(doc, department=department)

        # ================================================================
        # PART 1: Cover Page
        # ================================================================
        for _ in range(4):
            doc.add_paragraph("")

        grant_name = proposal.get("title", "Grant Proposal")
        if grant_context and grant_context.get("grant_name"):
            grant_name = grant_context["grant_name"]

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(grant_name)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = _BRAND_COLOR

        doc.add_paragraph("")

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle_para.add_run("Complete Application Package")
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        applicant_para = doc.add_paragraph()
        applicant_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        applicant_run = applicant_para.add_run("Applicant: City of Austin")
        applicant_run.font.size = Pt(12)
        applicant_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph("")
        hr_para = doc.add_paragraph()
        hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr_run = hr_para.add_run("_" * 60)
        hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        doc.add_page_break()

        # ================================================================
        # PART 2: Proposal Narrative
        # ================================================================
        doc.add_heading("PROPOSAL NARRATIVE", level=1)
        doc.add_paragraph("")

        sections = proposal.get("sections") or {}

        for section_key in _STANDARD_SECTIONS:
            section_data = sections.get(section_key)
            if not section_data:
                continue

            if isinstance(section_data, dict):
                content = section_data.get("content", "")
            elif isinstance(section_data, str):
                content = section_data
            else:
                continue

            if not content or not content.strip():
                continue

            doc.add_heading(svc._section_title(section_key), level=2)
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                stripped = para_text.strip()
                if stripped:
                    doc.add_paragraph(stripped)

        # Custom sections
        for section_key, section_data in sections.items():
            if section_key in _STANDARD_SECTIONS:
                continue

            if isinstance(section_data, dict):
                content = section_data.get("content", "")
            elif isinstance(section_data, str):
                content = section_data
            else:
                continue

            if not content or not content.strip():
                continue

            doc.add_heading(svc._section_title(section_key), level=2)
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                stripped = para_text.strip()
                if stripped:
                    doc.add_paragraph(stripped)

        # ================================================================
        # PART 3: Budget Detail
        # ================================================================
        doc.add_page_break()
        doc.add_heading("BUDGET DETAIL", level=1)

        if budget_items:
            items_by_category: dict[str, list[dict]] = {}
            for item in budget_items:
                cat = item.get("category", "other")
                items_by_category.setdefault(cat, []).append(item)

            for category in _FEDERAL_BUDGET_CATEGORIES:
                cat_items = items_by_category.get(category, [])
                if not cat_items:
                    continue

                doc.add_heading(svc._section_title(category), level=2)

                if category == "personnel":
                    headers = ["Role", "FTE", "Salary", "Months", "Total"]
                    widths = [2.0, 0.8, 1.2, 0.8, 1.2]
                    table = svc._add_table_with_header(doc, headers, widths)

                    for item in cat_items:
                        row = table.add_row()
                        row.cells[0].text = item.get(
                            "role", item.get("description", "")
                        )
                        row.cells[1].text = str(item.get("fte", ""))
                        row.cells[2].text = svc._format_currency(
                            item.get("annual_salary")
                        )
                        row.cells[3].text = str(item.get("months_on_project", ""))
                        row.cells[4].text = svc._format_currency(item.get("total_cost"))
                else:
                    headers = [
                        "Description",
                        "Qty",
                        "Unit Cost",
                        "Total",
                        "Justification",
                    ]
                    widths = [2.0, 0.6, 1.0, 1.0, 2.0]
                    table = svc._add_table_with_header(doc, headers, widths)

                    for item in cat_items:
                        row = table.add_row()
                        row.cells[0].text = item.get("description", "")
                        row.cells[1].text = str(item.get("quantity", ""))
                        row.cells[2].text = svc._format_currency(item.get("unit_cost"))
                        row.cells[3].text = svc._format_currency(item.get("total_cost"))
                        row.cells[4].text = item.get("justification", "") or ""

                doc.add_paragraph("")

            # Handle non-standard categories
            for category, cat_items in items_by_category.items():
                if category in _FEDERAL_BUDGET_CATEGORIES:
                    continue

                doc.add_heading(svc._section_title(category), level=2)
                headers = ["Description", "Qty", "Unit Cost", "Total", "Justification"]
                widths = [2.0, 0.6, 1.0, 1.0, 2.0]
                table = svc._add_table_with_header(doc, headers, widths)

                for item in cat_items:
                    row = table.add_row()
                    row.cells[0].text = item.get("description", "")
                    row.cells[1].text = str(item.get("quantity", ""))
                    row.cells[2].text = svc._format_currency(item.get("unit_cost"))
                    row.cells[3].text = svc._format_currency(item.get("total_cost"))
                    row.cells[4].text = item.get("justification", "") or ""

                doc.add_paragraph("")

            # Budget summary
            doc.add_heading("Budget Summary", level=2)
            s_headers = ["Item", "Amount"]
            s_widths = [4.0, 2.0]
            s_table = svc._add_table_with_header(doc, s_headers, s_widths)

            row = s_table.add_row()
            row.cells[0].text = "Direct Costs"
            row.cells[1].text = svc._format_currency(
                budget_calculations.get("direct_total", 0)
            )

            row = s_table.add_row()
            row.cells[0].text = "Indirect Costs"
            row.cells[1].text = svc._format_currency(
                budget_calculations.get("indirect_total", 0)
            )

            total_row = s_table.add_row()
            total_row.cells[0].text = "GRAND TOTAL"
            total_row.cells[1].text = svc._format_currency(
                budget_calculations.get("grand_total", 0)
            )
            for cell in total_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
        else:
            doc.add_paragraph("No budget line items have been entered.")

        # ================================================================
        # PART 4: Requirements Checklist
        # ================================================================
        doc.add_page_break()
        doc.add_heading("REQUIREMENTS CHECKLIST", level=1)

        if checklist_items:
            cl_headers = ["Description", "Status", "Mandatory"]
            cl_widths = [4.0, 1.5, 1.0]
            cl_table = svc._add_table_with_header(doc, cl_headers, cl_widths)

            completed_count = 0
            total_count = len(checklist_items)

            for item in checklist_items:
                row = cl_table.add_row()
                row.cells[0].text = item.get("description", "")
                is_completed = item.get("is_completed", False)
                row.cells[1].text = "Complete" if is_completed else "Incomplete"
                row.cells[2].text = "Yes" if item.get("is_mandatory", False) else "No"
                if is_completed:
                    completed_count += 1

            doc.add_paragraph("")
            summary_para = doc.add_paragraph()
            summary_run = summary_para.add_run(
                f"Completion: {completed_count}/{total_count} items complete"
            )
            summary_run.bold = True
        else:
            doc.add_paragraph("No checklist items have been defined.")

        # ---- Save to bytes ----
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # ------------------------------------------------------------------
    # Shared DOCX helpers (program summary / project plan)
    # ------------------------------------------------------------------

    @staticmethod
    def _add_cover_page(
        doc: Document,
        title: str,
        subtitle: str | None = None,
        author: str | None = None,
    ) -> None:
        """Add a standard cover page: padding, title, subtitle, author, date, HR, page break.

        Args:
            doc: python-docx Document to append to.
            title: Document title (bold, 18pt, brand color, centered).
            subtitle: Optional subtitle (14pt, gray, centered).
            author: Optional "Prepared by ..." line (12pt, gray, centered).
        """
        # Top padding
        for _ in range(4):
            doc.add_paragraph("")

        # Title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = _BRAND_COLOR

        doc.add_paragraph("")

        # Subtitle
        if subtitle:
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # Author
        if author:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_para.add_run(f"Prepared by {author}")
            author_run.font.size = Pt(12)
            author_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(datetime.now(timezone.utc).strftime("%B %d, %Y"))
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Horizontal rule
        doc.add_paragraph("")
        hr_para = doc.add_paragraph()
        hr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr_run = hr_para.add_run("_" * 60)
        hr_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

        doc.add_page_break()

    @staticmethod
    def _render_paragraph_section(
        doc: Document,
        heading: str,
        text: str,
    ) -> None:
        """Add a heading followed by paragraph text (split on double-newlines).

        Args:
            doc: python-docx Document to append to.
            heading: Section heading text (rendered at level 1).
            text: Body text; if empty/whitespace-only the section is skipped.
        """
        text = str(text).strip()
        if not text:
            return

        doc.add_heading(heading, level=1)
        paragraphs = text.split("\n\n")
        for para_text in paragraphs:
            stripped = para_text.strip()
            if stripped:
                doc.add_paragraph(stripped)

    @staticmethod
    def _render_bullet_section(
        doc: Document,
        heading: str,
        items: list,
    ) -> None:
        """Add a heading followed by a bulleted list of items.

        Args:
            doc: python-docx Document to append to.
            heading: Section heading text (rendered at level 1).
            items: List of items to render as "List Bullet" paragraphs.
                Empty/whitespace-only items are silently skipped.
        """
        if not items:
            return

        cleaned = [str(i).strip() for i in items if str(i).strip()]
        if not cleaned:
            return

        doc.add_heading(heading, level=1)
        for item_text in cleaned:
            doc.add_paragraph(item_text, style="List Bullet")

    # ------------------------------------------------------------------
    # Program Summary DOCX
    # ------------------------------------------------------------------

    @staticmethod
    def generate_program_summary_docx(
        summary_data: dict,
        profile_data: dict | None = None,
    ) -> io.BytesIO:
        """Generate a DOCX program summary document.

        Args:
            summary_data: ProgramSummary fields (program_name, department,
                problem_statement, program_description, target_population,
                key_needs, estimated_budget, team_overview, timeline_overview,
                strategic_alignment)
            profile_data: Optional user profile context for author info

        Returns:
            BytesIO buffer containing the DOCX file
        """
        doc = Document()
        svc = DocxExportService

        svc._set_margins(doc, inches=1.0)
        svc._set_default_font(doc, name="Times New Roman", size=12)

        department = (
            summary_data.get("department", "City of Austin") or "City of Austin"
        )
        svc._add_page_headers_and_footers(doc, department=department)

        # ---- Cover / Title Area ----
        program_name = summary_data.get("program_name", "")

        # Build author string from profile_data
        author_str = None
        if profile_data:
            display_name = profile_data.get("display_name", "")
            if display_name:
                author_parts = [display_name]
                prof_dept = profile_data.get("department", "")
                if prof_dept:
                    author_parts.append(prof_dept)
                author_str = ", ".join(author_parts)

        svc._add_cover_page(
            doc,
            title="Program Summary",
            subtitle=program_name or None,
            author=author_str,
        )

        # ---- Sections (only include if data is non-empty) ----
        _summary_sections = [
            ("Problem Statement", "problem_statement"),
            ("Program Description", "program_description"),
            ("Target Population", "target_population"),
            ("Key Needs", "key_needs"),
            ("Budget Overview", "estimated_budget"),
            ("Team Overview", "team_overview"),
            ("Timeline", "timeline_overview"),
            ("Strategic Alignment", "strategic_alignment"),
        ]

        for heading, key in _summary_sections:
            value = summary_data.get(key)
            if not value:
                continue

            # Key Needs is rendered as a bulleted list
            if key == "key_needs":
                if isinstance(value, list) and len(value) > 0:
                    svc._render_bullet_section(doc, heading, value)
                continue

            # Everything else is paragraph text
            svc._render_paragraph_section(doc, heading, str(value))

        # ---- Save to BytesIO ----
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # ------------------------------------------------------------------
    # Project Plan DOCX
    # ------------------------------------------------------------------

    @staticmethod
    def generate_project_plan_docx(
        plan_data: dict,
        grant_context: dict | None = None,
        profile_data: dict | None = None,
    ) -> io.BytesIO:
        """Generate a DOCX project plan document.

        Args:
            plan_data: PlanData fields (program_overview, staffing_plan, budget,
                timeline, deliverables, metrics, partnerships)
            grant_context: Optional grant context for header info
            profile_data: Optional user profile context

        Returns:
            BytesIO buffer containing the DOCX file
        """
        doc = Document()
        svc = DocxExportService

        svc._set_margins(doc, inches=1.0)
        svc._set_default_font(doc, name="Times New Roman", size=12)

        department = "City of Austin"
        if grant_context and grant_context.get("department"):
            department = grant_context["department"]
        elif profile_data and profile_data.get("department"):
            department = profile_data["department"]

        svc._add_page_headers_and_footers(doc, department=department)

        # ---- Cover / Title Area ----
        # Subtitle: first line of program_overview or profile_data program_name
        subtitle_text = ""
        overview = plan_data.get("program_overview", "")
        if overview:
            subtitle_text = str(overview).split("\n")[0].strip()
        if not subtitle_text and profile_data:
            subtitle_text = profile_data.get("program_name", "")

        if subtitle_text and len(subtitle_text) > 120:
            subtitle_text = subtitle_text[:117] + "..."

        # Build author string from profile_data
        author_str = None
        if profile_data:
            display_name = profile_data.get("display_name", "")
            if display_name:
                author_parts = [display_name]
                prof_dept = profile_data.get("department", "")
                if prof_dept:
                    author_parts.append(prof_dept)
                author_str = ", ".join(author_parts)

        svc._add_cover_page(
            doc,
            title="Project Plan",
            subtitle=subtitle_text or None,
            author=author_str,
        )

        # ---- Grant Opportunity (if grant_context provided) ----
        if grant_context:
            doc.add_heading("Grant Opportunity", level=1)

            gc_fields = [
                (
                    "Grant Name",
                    grant_context.get("grant_name") or grant_context.get("name"),
                ),
                ("Grantor", grant_context.get("grantor")),
                ("Deadline", grant_context.get("deadline")),
            ]

            # Funding range
            funding_min = grant_context.get("funding_amount_min")
            funding_max = grant_context.get("funding_amount_max")
            if funding_min and funding_max:
                gc_fields.append(
                    (
                        "Funding Range",
                        f"${float(funding_min):,.0f} \u2013 ${float(funding_max):,.0f}",
                    )
                )
            elif funding_max:
                gc_fields.append(
                    ("Funding Amount", f"Up to ${float(funding_max):,.0f}")
                )
            elif funding_min:
                gc_fields.append(("Funding Amount", f"From ${float(funding_min):,.0f}"))

            for label, value in gc_fields:
                if value:
                    para = doc.add_paragraph()
                    label_run = para.add_run(f"{label}: ")
                    label_run.bold = True
                    para.add_run(str(value))

            doc.add_paragraph("")

        # ---- Program Narrative ----
        narrative = plan_data.get("program_overview", "")
        svc._render_paragraph_section(doc, "Program Narrative", narrative)

        # ---- Staffing Plan Table ----
        staffing = plan_data.get("staffing_plan")
        if staffing and isinstance(staffing, list) and len(staffing) > 0:
            doc.add_heading("Staffing Plan", level=1)
            headers = ["Role", "FTE", "Salary Estimate", "Responsibilities"]
            widths = [1.8, 0.7, 1.3, 3.0]
            table = svc._add_table_with_header(doc, headers, widths)

            for entry in staffing:
                if not isinstance(entry, dict):
                    continue
                row = table.add_row()
                row.cells[0].text = str(entry.get("role", ""))
                row.cells[1].text = str(entry.get("fte", ""))
                row.cells[2].text = svc._format_currency(
                    entry.get("salary") or entry.get("salary_estimate")
                )
                row.cells[3].text = str(entry.get("responsibilities", ""))

            doc.add_paragraph("")

        # ---- Budget Breakdown Table ----
        budget = plan_data.get("budget")
        if budget and isinstance(budget, list) and len(budget) > 0:
            doc.add_heading("Budget Breakdown", level=1)
            headers = ["Category", "Amount ($)", "Justification"]
            widths = [2.0, 1.3, 3.5]
            table = svc._add_table_with_header(doc, headers, widths)

            total_amount = 0.0
            for entry in budget:
                if not isinstance(entry, dict):
                    continue
                row = table.add_row()
                row.cells[0].text = str(entry.get("category", ""))
                amount = entry.get("amount", 0)
                amount_val = float(amount) if amount else 0
                total_amount += amount_val
                row.cells[1].text = svc._format_currency(amount_val)
                row.cells[2].text = str(entry.get("justification", ""))

            # Total row (bold)
            total_row = table.add_row()
            total_row.cells[0].text = "TOTAL"
            total_row.cells[1].text = svc._format_currency(total_amount)
            total_row.cells[2].text = ""
            for cell in total_row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            doc.add_paragraph("")

        # ---- Timeline ----
        timeline = plan_data.get("timeline")
        if timeline and isinstance(timeline, list) and len(timeline) > 0:
            doc.add_heading("Timeline", level=1)

            for phase in timeline:
                if not isinstance(phase, dict):
                    continue

                phase_name = phase.get("phase") or phase.get("name", "Phase")
                start = phase.get("start", "")
                end = phase.get("end", "")
                date_range = ""
                if start and end:
                    date_range = f" ({start} \u2013 {end})"
                elif start:
                    date_range = f" (from {start})"
                elif end:
                    date_range = f" (until {end})"

                doc.add_heading(f"{phase_name}{date_range}", level=2)

                milestones = phase.get("milestones", [])
                if isinstance(milestones, list):
                    for ms in milestones:
                        ms_text = str(ms).strip()
                        if ms_text:
                            doc.add_paragraph(ms_text, style="List Bullet")

                # Also handle a description field if present
                desc = phase.get("description", "")
                if desc and str(desc).strip():
                    doc.add_paragraph(str(desc).strip())

        # ---- Deliverables ----
        deliverables = plan_data.get("deliverables")
        if deliverables and isinstance(deliverables, list) and len(deliverables) > 0:
            svc._render_bullet_section(doc, "Deliverables", deliverables)

        # ---- Success Metrics Table ----
        metrics = plan_data.get("metrics")
        if metrics and isinstance(metrics, list) and len(metrics) > 0:
            doc.add_heading("Success Metrics", level=1)
            headers = ["Metric", "Target", "How Measured"]
            widths = [2.5, 1.5, 3.0]
            table = svc._add_table_with_header(doc, headers, widths)

            for entry in metrics:
                if not isinstance(entry, dict):
                    continue
                row = table.add_row()
                row.cells[0].text = str(entry.get("metric", ""))
                row.cells[1].text = str(entry.get("target", ""))
                row.cells[2].text = str(
                    entry.get("how_measured") or entry.get("measurement", "")
                )

            doc.add_paragraph("")

        # ---- Partnerships (if present) ----
        partnerships = plan_data.get("partnerships")
        if partnerships:
            if isinstance(partnerships, list) and len(partnerships) > 0:
                svc._render_bullet_section(doc, "Partnerships", partnerships)
            elif isinstance(partnerships, str) and partnerships.strip():
                svc._render_paragraph_section(doc, "Partnerships", partnerships)

        # ---- Save to BytesIO ----
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

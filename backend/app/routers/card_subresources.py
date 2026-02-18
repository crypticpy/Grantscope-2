"""Card sub-resource router -- sources, timeline, history, related, follow, notes, assets, velocity, documents.

Migrated from Supabase PostgREST to SQLAlchemy 2.0 async.
"""

import io
import logging
import re
import uuid as _uuid
from datetime import date, datetime, timedelta, timezone
from decimal import Decimal
from typing import Dict, Any, List, Optional

from fastapi import (
    APIRouter,
    Depends,
    File,
    Form,
    HTTPException,
    Query,
    UploadFile,
    status,
)
from pydantic import BaseModel
from sqlalchemy import or_, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.deps import get_db, get_current_user_hardcoded, _safe_error
from app.models.db.card import Card
from app.models.db.card_extras import (
    CardFollow,
    CardNote,
    CardTimeline,
    CardScoreHistory,
    CardRelationship,
    Entity,
    UserSignalPreference,
)
from app.models.db.source import Source
from app.models.db.research import ResearchTask
from app.models.db.brief import ExecutiveBrief
from app.models.db.workstream import Workstream, WorkstreamCard
from app.models.history import (
    ScoreHistory,
    ScoreHistoryResponse,
    StageHistory,
    StageHistoryList,
    RelatedCard,
    RelatedCardsList,
)
from app.models.workstream import Note, NoteCreate
from app.models.assets import CardAsset, CardAssetsResponse

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1", tags=["card-subresources"])

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_SKIP_COLS = {"embedding", "search_vector", "content_embedding"}


def _row_to_dict(obj, skip_cols: set | None = None) -> dict:
    """Convert an ORM object to a plain dict, serialising special types."""
    skip = skip_cols or _SKIP_COLS
    result: dict[str, Any] = {}
    for col in obj.__table__.columns:
        if col.name in skip:
            continue
        value = getattr(obj, col.name, None)
        if isinstance(value, _uuid.UUID):
            result[col.name] = str(value)
        elif isinstance(value, (datetime, date)):
            result[col.name] = value.isoformat()
        elif isinstance(value, Decimal):
            result[col.name] = float(value)
        else:
            result[col.name] = value
    return result


# ============================================================================
# Entity models
# ============================================================================


class EntityItem(BaseModel):
    id: str
    name: str
    entity_type: str
    context: Optional[str] = None
    source_id: Optional[str] = None
    canonical_name: Optional[str] = None
    created_at: str


class EntityListResponse(BaseModel):
    entities: List[EntityItem]
    total_count: int
    card_id: str


# ============================================================================
# Card relationships / sources / timeline
# ============================================================================


@router.get("/cards/{card_id}/sources")
async def get_card_sources(
    card_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Get sources for a card"""
    try:
        result = await db.execute(
            select(Source)
            .where(Source.card_id == card_id)
            .order_by(Source.relevance_score.desc().nulls_last())
        )
        sources = result.scalars().all()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card sources retrieval", e),
        ) from e
    return [_row_to_dict(s) for s in sources]


@router.get("/cards/{card_id}/timeline")
async def get_card_timeline(
    card_id: str,
    db: AsyncSession = Depends(get_db),
):
    """Get timeline for a card"""
    try:
        result = await db.execute(
            select(CardTimeline)
            .where(CardTimeline.card_id == card_id)
            .order_by(CardTimeline.created_at.desc())
        )
        events = result.scalars().all()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card timeline retrieval", e),
        ) from e
    return [_row_to_dict(ev) for ev in events]


@router.get("/cards/{card_id}/entities", response_model=EntityListResponse)
async def get_card_entities(
    card_id: str,
    entity_type: Optional[str] = None,
    limit: int = 50,
    db: AsyncSession = Depends(get_db),
):
    """
    Get entities extracted from a card's sources.

    Returns entities (technologies, organizations, concepts, people, locations)
    associated with the given card, optionally filtered by entity type.

    Args:
        card_id: UUID of the card to get entities for
        entity_type: Optional filter by entity type (technology, organization,
                     concept, person, location)
        limit: Maximum number of entities to return (default: 50)

    Returns:
        EntityListResponse with list of entities and metadata
    """
    try:
        # First verify the card exists
        card_result = await db.execute(select(Card.id).where(Card.id == card_id))
        if card_result.scalar_one_or_none() is None:
            raise HTTPException(status_code=404, detail="Card not found")

        # Build query for entities
        stmt = select(Entity).where(Entity.card_id == card_id)

        # Apply optional entity_type filter
        if entity_type:
            stmt = stmt.where(Entity.entity_type == entity_type)

        # Execute query ordered by name, with limit
        stmt = stmt.order_by(Entity.name).limit(limit)
        result = await db.execute(stmt)
        rows = result.scalars().all()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card entities retrieval", e),
        ) from e

    # Convert to EntityItem models
    entities = [
        EntityItem(
            id=str(ent.id),
            name=ent.name,
            entity_type=ent.entity_type,
            context=ent.context,
            source_id=str(ent.source_id) if ent.source_id else None,
            canonical_name=ent.canonical_name,
            created_at=ent.created_at.isoformat() if ent.created_at else "",
        )
        for ent in rows
    ]

    return EntityListResponse(
        entities=entities,
        total_count=len(entities),
        card_id=card_id,
    )


@router.get("/cards/{card_id}/score-history", response_model=ScoreHistoryResponse)
async def get_card_score_history(
    card_id: str,
    start_date: Optional[datetime] = None,
    end_date: Optional[datetime] = None,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """
    Get historical score data for a card to enable trend visualization.

    Returns a list of score snapshots ordered by recorded_at (most recent first),
    containing all 7 score dimensions (maturity, velocity, novelty, impact,
    relevance, risk, opportunity) for each timestamp.

    Args:
        card_id: UUID of the card to get score history for
        start_date: Optional filter to get records from this date onwards
        end_date: Optional filter to get records up to this date

    Returns:
        ScoreHistoryResponse with list of ScoreHistory records and metadata
    """
    try:
        # First verify the card exists
        card_result = await db.execute(select(Card.id).where(Card.id == card_id))
        if card_result.scalar_one_or_none() is None:
            raise HTTPException(status_code=404, detail="Card not found")

        # Build query for score history
        stmt = select(CardScoreHistory).where(CardScoreHistory.card_id == card_id)

        # Apply date filters if provided
        if start_date:
            stmt = stmt.where(CardScoreHistory.recorded_at >= start_date)
        if end_date:
            stmt = stmt.where(CardScoreHistory.recorded_at <= end_date)

        # Execute query ordered by recorded_at descending
        stmt = stmt.order_by(CardScoreHistory.recorded_at.desc())
        result = await db.execute(stmt)
        rows = result.scalars().all()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card score history retrieval", e),
        ) from e

    # Convert to ScoreHistory models
    history_records = [
        ScoreHistory(
            id=str(row.id),
            card_id=str(row.card_id),
            recorded_at=row.recorded_at,
            maturity_score=row.maturity_score,
            velocity_score=row.velocity_score,
            novelty_score=row.novelty_score,
            impact_score=row.impact_score,
            relevance_score=row.relevance_score,
            risk_score=row.risk_score,
            opportunity_score=row.opportunity_score,
        )
        for row in rows
    ]

    return ScoreHistoryResponse(
        history=history_records,
        card_id=card_id,
        total_count=len(history_records),
        start_date=start_date,
        end_date=end_date,
    )


@router.get("/cards/{card_id}/stage-history", response_model=StageHistoryList)
async def get_card_stage_history(
    card_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """
    Get maturity stage transition history for a card.

    Returns a list of stage transitions ordered by changed_at (most recent first),
    tracking maturity stage progression through stages 1-8 and horizon shifts
    (H3 -> H2 -> H1).

    The data is sourced from the card_timeline table, filtered to only include
    'stage_changed' event types.

    Args:
        card_id: UUID of the card to get stage history for

    Returns:
        StageHistoryList with stage transition records and metadata
    """
    try:
        # First verify the card exists
        card_result = await db.execute(select(Card.id).where(Card.id == card_id))
        if card_result.scalar_one_or_none() is None:
            raise HTTPException(status_code=404, detail="Card not found")

        # Query card_timeline for stage change events
        result = await db.execute(
            select(CardTimeline)
            .where(
                CardTimeline.card_id == card_id,
                CardTimeline.event_type == "stage_changed",
            )
            .order_by(CardTimeline.created_at.desc())
        )
        rows = result.scalars().all()
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card stage history retrieval", e),
        ) from e

    # Convert to StageHistory models, mapping created_at to changed_at
    history_records = []
    for row in rows:
        if row.new_stage_id is None:
            continue
        history_records.append(
            StageHistory(
                id=str(row.id),
                card_id=str(row.card_id),
                changed_at=row.created_at,  # Map created_at to changed_at
                old_stage_id=row.old_stage_id,
                new_stage_id=row.new_stage_id,
                old_horizon=row.old_horizon,
                new_horizon=row.new_horizon or "H3",  # Default to H3 if not set
                trigger=row.trigger,
                reason=row.reason,
            )
        )

    return StageHistoryList(
        history=history_records, total_count=len(history_records), card_id=card_id
    )


@router.get("/cards/{card_id}/related", response_model=RelatedCardsList)
async def get_related_cards(
    card_id: str,
    limit: int = 20,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """
    Get cards related to the specified card for concept network visualization.

    Returns cards connected to the source card through the card_relationships table,
    including relationship metadata (type and strength) for edge visualization.
    Relationships are bidirectional - cards appear whether they are source or target.

    Args:
        card_id: UUID of the source card to get relationships for
        limit: Maximum number of related cards to return (default: 20)

    Returns:
        RelatedCardsList with related card details and relationship metadata
    """
    try:
        # First verify the card exists
        card_result = await db.execute(select(Card.id).where(Card.id == card_id))
        if card_result.scalar_one_or_none() is None:
            raise HTTPException(status_code=404, detail="Card not found")

        # Query relationships where this card is either source or target
        result = await db.execute(
            select(CardRelationship).where(
                or_(
                    CardRelationship.source_card_id == card_id,
                    CardRelationship.target_card_id == card_id,
                )
            )
        )
        all_relationships = list(result.scalars().all())
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card relationships retrieval", e),
        ) from e

    # If no relationships found, return empty list
    if not all_relationships:
        return RelatedCardsList(related_cards=[], total_count=0, source_card_id=card_id)

    # Get the related card IDs (the "other" card in each relationship)
    related_card_ids = set()
    for rel in all_relationships:
        src = str(rel.source_card_id)
        tgt = str(rel.target_card_id)
        if src == card_id:
            related_card_ids.add(tgt)
        else:
            related_card_ids.add(src)

    try:
        # Fetch full card details for all related cards
        cards_result = await db.execute(
            select(Card).where(Card.id.in_(list(related_card_ids)))
        )
        related_card_objs = cards_result.scalars().all()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("related cards retrieval", e),
        ) from e

    # Create a lookup map for cards
    cards_map = {str(c.id): c for c in related_card_objs}

    # Build the related cards list with relationship context
    related_cards: list[RelatedCard] = []
    for rel in all_relationships:
        src = str(rel.source_card_id)
        tgt = str(rel.target_card_id)
        # Determine which card is the "related" one (not the source card_id)
        related_id = tgt if src == card_id else src

        card_data = cards_map.get(related_id)
        if card_data is not None:
            related_cards.append(
                RelatedCard(
                    id=str(card_data.id),
                    name=card_data.name,
                    slug=card_data.slug,
                    summary=card_data.summary,
                    pillar_id=card_data.pillar_id,
                    stage_id=card_data.stage_id,
                    horizon=card_data.horizon,
                    relationship_type=rel.relationship_type,
                    relationship_strength=(
                        float(rel.strength) if rel.strength is not None else None
                    ),
                    relationship_id=str(rel.id),
                )
            )

    # Limit the results to the specified limit
    related_cards = related_cards[:limit]

    return RelatedCardsList(
        related_cards=related_cards,
        total_count=len(related_cards),
        source_card_id=card_id,
    )


# ============================================================================
# Follow / unfollow
# ============================================================================


@router.post("/cards/{card_id}/follow")
async def follow_card(
    card_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Follow a card"""
    try:
        follow = CardFollow(
            user_id=_uuid.UUID(current_user["id"]),
            card_id=_uuid.UUID(card_id),
        )
        db.add(follow)
        await db.flush()
    except Exception as e:
        logger.warning("Failed to follow card %s: %s", card_id, e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("follow card", e),
        ) from e

    try:
        from app.signal_quality import update_signal_quality_score

        await update_signal_quality_score(db, card_id)
    except Exception as e:
        logger.warning(f"Failed to update signal quality score for {card_id}: {e}")

    return {"status": "followed"}


@router.delete("/cards/{card_id}/follow")
async def unfollow_card(
    card_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Unfollow a card"""
    try:
        result = await db.execute(
            select(CardFollow).where(
                CardFollow.user_id == current_user["id"],
                CardFollow.card_id == card_id,
            )
        )
        follow = result.scalar_one_or_none()
        if follow:
            await db.delete(follow)
            await db.flush()
    except Exception as e:
        logger.warning("Failed to unfollow card %s: %s", card_id, e)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("unfollow card", e),
        ) from e

    try:
        from app.signal_quality import update_signal_quality_score

        await update_signal_quality_score(db, card_id)
    except Exception as e:
        logger.warning(f"Failed to update signal quality score for {card_id}: {e}")

    return {"status": "unfollowed"}


# ============================================================================
# Following / My Signals
# ============================================================================


@router.get("/me/following")
async def get_following_cards(
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Get cards followed by current user"""
    try:
        user_id = current_user["id"]

        # Get all follows for the user
        follows_result = await db.execute(
            select(CardFollow).where(CardFollow.user_id == user_id)
        )
        follows = follows_result.scalars().all()

        if not follows:
            return []

        # Get the card IDs
        card_ids = [str(f.card_id) for f in follows]

        # Fetch full card data
        cards_result = await db.execute(select(Card).where(Card.id.in_(card_ids)))
        cards = cards_result.scalars().all()
        cards_map = {str(c.id): c for c in cards}

        # Build response: follow record with nested card data
        response_data = []
        for f in follows:
            follow_dict = _row_to_dict(f)
            card = cards_map.get(str(f.card_id))
            if card:
                follow_dict["cards"] = _row_to_dict(card)
            response_data.append(follow_dict)

        return response_data
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("following cards retrieval", e),
        ) from e


@router.get("/me/signals")
async def get_my_signals(
    group_by: Optional[str] = Query(
        None, description="Group by: pillar, horizon, workstream"
    ),
    sort_by: str = Query(
        "updated", description="Sort: updated, followed, quality, name"
    ),
    search: Optional[str] = Query(None, description="Search term"),
    pillar: Optional[str] = Query(None, description="Filter by pillar"),
    horizon: Optional[str] = Query(None, description="Filter by horizon"),
    source: Optional[str] = Query(
        None, description="Filter by: followed, created, workstream"
    ),
    quality_min: Optional[int] = Query(None, ge=0, le=100),
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Get user's personal intelligence hub: followed, created, and workstream signals."""
    user_id = current_user["id"]

    try:
        # 1. Get followed card IDs
        follows_result = await db.execute(
            select(CardFollow).where(CardFollow.user_id == user_id)
        )
        follows = follows_result.scalars().all()
        followed_map = {
            str(f.card_id): {
                "card_id": str(f.card_id),
                "created_at": f.created_at.isoformat() if f.created_at else None,
                "priority": f.priority,
                "notes": f.notes,
            }
            for f in follows
        }
        followed_ids = list(followed_map.keys())

        # 2. Get user-created card IDs
        created_result = await db.execute(
            select(Card.id).where(Card.created_by == user_id, Card.status == "active")
        )
        created_ids = [str(row[0]) for row in created_result.all()]

        # 3. Get cards in user's workstreams
        ws_result = await db.execute(
            select(Workstream).where(Workstream.user_id == user_id)
        )
        workstreams_objs = ws_result.scalars().all()
        workstreams = [{"id": str(ws.id), "name": ws.name} for ws in workstreams_objs]
        ws_ids = [str(ws.id) for ws in workstreams_objs]

        ws_card_ids: list[str] = []
        ws_card_map: Dict[str, List[str]] = {}  # card_id -> list of workstream names
        if ws_ids:
            wc_result = await db.execute(
                select(WorkstreamCard).where(WorkstreamCard.workstream_id.in_(ws_ids))
            )
            wc_rows = wc_result.scalars().all()
            ws_name_map = {ws["id"]: ws["name"] for ws in workstreams}
            for wc in wc_rows:
                cid = str(wc.card_id)
                ws_card_ids.append(cid)
                if cid not in ws_card_map:
                    ws_card_map[cid] = []
                ws_card_map[cid].append(
                    ws_name_map.get(str(wc.workstream_id), "Unknown")
                )

        # 4. Union unique card IDs, applying source filter if specified
        if source == "followed":
            all_ids = list(set(followed_ids))
        elif source == "created":
            all_ids = list(set(created_ids))
        elif source == "workstream":
            all_ids = list(set(ws_card_ids))
        else:
            all_ids = list(set(followed_ids + created_ids + ws_card_ids))

        if not all_ids:
            return {
                "signals": [],
                "stats": {
                    "total": 0,
                    "followed_count": 0,
                    "created_count": 0,
                    "workstream_count": len(workstreams),
                    "updates_this_week": 0,
                    "needs_research": 0,
                },
                "workstreams": workstreams,
            }

        # 5. Fetch full card data for all IDs
        cards_stmt = select(Card).where(Card.id.in_(all_ids), Card.status == "active")

        if search:
            safe_search = re.sub(r"[,.()\[\]]", "", search)
            cards_stmt = cards_stmt.where(
                or_(
                    Card.name.ilike(f"%{safe_search}%"),
                    Card.summary.ilike(f"%{safe_search}%"),
                )
            )
        if pillar:
            cards_stmt = cards_stmt.where(Card.pillar_id == pillar)
        if horizon:
            cards_stmt = cards_stmt.where(Card.horizon == horizon)
        if quality_min is not None and quality_min > 0:
            cards_stmt = cards_stmt.where(Card.signal_quality_score >= quality_min)

        cards_result = await db.execute(cards_stmt)
        cards = cards_result.scalars().all()

        # 6. Get user signal preferences (pins) -- gracefully degrade if table missing
        prefs_map: dict[str, dict] = {}
        try:
            prefs_result = await db.execute(
                select(UserSignalPreference).where(
                    UserSignalPreference.user_id == user_id
                )
            )
            prefs = prefs_result.scalars().all()
            prefs_map = {str(p.card_id): _row_to_dict(p) for p in prefs}
        except Exception:
            logger.warning(
                "user_signal_preferences table may not exist; skipping pin data"
            )

        # 7. Enrich cards with personal metadata
        one_week_ago = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
        enriched = []
        for card in cards:
            card_dict = _row_to_dict(card)
            cid = str(card.id)
            pref = prefs_map.get(cid, {})
            follow_data = followed_map.get(cid)
            card_dict.update(
                {
                    "is_followed": cid in followed_ids,
                    "is_created": cid in created_ids,
                    "is_pinned": pref.get("is_pinned", False),
                    "personal_notes": pref.get("notes"),
                    "follow_priority": (
                        follow_data.get("priority") if follow_data else None
                    ),
                    "followed_at": (
                        follow_data.get("created_at") if follow_data else None
                    ),
                    "workstream_names": ws_card_map.get(cid, []),
                }
            )
            enriched.append(card_dict)

        # 8. Sort
        if sort_by == "quality":
            enriched.sort(
                key=lambda c: c.get("signal_quality_score") or 0, reverse=True
            )
        elif sort_by == "followed":
            enriched.sort(key=lambda c: c.get("followed_at") or "", reverse=True)
        elif sort_by == "name":
            enriched.sort(key=lambda c: c.get("name", "").lower())
        else:  # default: updated
            enriched.sort(
                key=lambda c: c.get("updated_at") or c.get("created_at") or "",
                reverse=True,
            )

        # Pinned first
        enriched.sort(key=lambda c: 0 if c.get("is_pinned") else 1)

        # 9. Stats
        updates_this_week = sum(
            bool((c.get("updated_at") or "") >= one_week_ago) for c in enriched
        )
        needs_research = sum(
            bool((c.get("signal_quality_score") or 0) < 30) for c in enriched
        )

        return {
            "signals": enriched,
            "stats": {
                "total": len(enriched),
                "followed_count": sum(bool(c.get("is_followed")) for c in enriched),
                "created_count": sum(bool(c.get("is_created")) for c in enriched),
                "workstream_count": len(workstreams),
                "updates_this_week": updates_this_week,
                "needs_research": needs_research,
            },
            "workstreams": workstreams,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("my signals retrieval", e),
        ) from e


# ============================================================================
# Pin signal
# ============================================================================


@router.post("/me/signals/{card_id}/pin")
async def pin_signal(
    card_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Pin/unpin a signal in the user's personal hub."""
    user_id = current_user["id"]

    try:
        # Check if preference exists
        result = await db.execute(
            select(UserSignalPreference).where(
                UserSignalPreference.user_id == user_id,
                UserSignalPreference.card_id == card_id,
            )
        )
        existing = result.scalar_one_or_none()

        if existing:
            # Toggle pin
            new_val = not (existing.is_pinned or False)
            existing.is_pinned = new_val
            existing.updated_at = datetime.now(timezone.utc)
            await db.flush()
            return {"is_pinned": new_val}
        else:
            # Create with pinned=True
            pref = UserSignalPreference(
                user_id=_uuid.UUID(user_id),
                card_id=_uuid.UUID(card_id),
                is_pinned=True,
            )
            db.add(pref)
            await db.flush()
            return {"is_pinned": True}
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("signal pin toggle", e),
        ) from e


# ============================================================================
# Notes
# ============================================================================


@router.get("/cards/{card_id}/notes")
async def get_card_notes(
    card_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Get notes for a card"""
    try:
        result = await db.execute(
            select(CardNote)
            .where(
                CardNote.card_id == card_id,
                or_(
                    CardNote.user_id == current_user["id"],
                    CardNote.is_private == False,  # noqa: E712
                ),
            )
            .order_by(CardNote.created_at.desc())
        )
        notes = result.scalars().all()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card notes retrieval", e),
        ) from e

    return [
        Note(
            id=str(n.id),
            content=n.content,
            is_private=n.is_private or False,
            created_at=n.created_at,
        )
        for n in notes
    ]


@router.post("/cards/{card_id}/notes")
async def create_note(
    card_id: str,
    note_data: NoteCreate,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Create note for a card"""
    try:
        note = CardNote(
            user_id=_uuid.UUID(current_user["id"]),
            card_id=_uuid.UUID(card_id),
            content=note_data.content,
            is_private=note_data.is_private,
            created_at=datetime.now(timezone.utc),
        )
        db.add(note)
        await db.flush()
        await db.refresh(note)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=_safe_error("note creation", e),
        ) from e

    return Note(
        id=str(note.id),
        content=note.content,
        is_private=note.is_private or False,
        created_at=note.created_at,
    )


# ============================================================================
# Assets
# ============================================================================


@router.get("/cards/{card_id}/assets", response_model=CardAssetsResponse)
async def get_card_assets(
    card_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """
    Get all generated assets for a card.

    Returns a list of all briefs, research reports, and exports
    associated with the card across all workstreams.

    Args:
        card_id: UUID of the card
        current_user: Authenticated user (injected)

    Returns:
        CardAssetsResponse with list of assets

    Raises:
        HTTPException 404: Card not found
    """
    try:
        # Verify card exists
        card_result = await db.execute(
            select(Card.id, Card.name).where(Card.id == card_id)
        )
        card_row = card_result.one_or_none()
        if not card_row:
            raise HTTPException(status_code=404, detail="Card not found")

        assets: list[CardAsset] = []

        # 1. Fetch executive briefs for this card
        briefs_result = await db.execute(
            select(ExecutiveBrief)
            .where(ExecutiveBrief.card_id == card_id)
            .order_by(ExecutiveBrief.created_at.desc())
        )
        briefs = briefs_result.scalars().all()

        for brief in briefs:
            # Map status
            brief_status = (
                "ready" if brief.status == "completed" else brief.status or "ready"
            )
            if brief_status == "generating":
                brief_status = "generating"
            elif brief_status in ("pending", "failed"):
                brief_status = "failed" if brief_status == "failed" else "ready"

            title = f"Executive Brief v{brief.version or 1}"

            created_at_val = brief.generated_at or brief.created_at

            assets.append(
                CardAsset(
                    id=str(brief.id),
                    type="brief",
                    title=title,
                    created_at=(created_at_val.isoformat() if created_at_val else ""),
                    version=brief.version or 1,
                    ai_generated=True,
                    ai_model=brief.model_used,
                    status=brief_status,
                    metadata={
                        "summary_preview": (
                            brief.summary[:200] if brief.summary else None
                        )
                    },
                )
            )

        # 2. Fetch research tasks (deep research reports)
        research_result = await db.execute(
            select(ResearchTask)
            .where(ResearchTask.card_id == card_id)
            .order_by(ResearchTask.created_at.desc())
        )
        tasks = research_result.scalars().all()

        for task in tasks:
            # Only include completed or failed tasks as assets
            if task.status not in ("completed", "failed"):
                continue

            task_type = task.task_type or "research"
            asset_type = "research"
            if task_type == "deep_research":
                title = "Strategic Intelligence Report"
            elif task_type == "update":
                title = "Quick Update Report"
            else:
                title = f"{task_type.replace('_', ' ').title()} Report"

            result_summary = task.result_summary or {}

            created_at_val = task.completed_at or task.created_at

            assets.append(
                CardAsset(
                    id=str(task.id),
                    type=asset_type,
                    title=title,
                    created_at=(created_at_val.isoformat() if created_at_val else ""),
                    ai_generated=True,
                    status="ready" if task.status == "completed" else "failed",
                    metadata={
                        "task_type": task_type,
                        "sources_found": result_summary.get("sources_found"),
                        "sources_added": result_summary.get("sources_added"),
                    },
                )
            )

        # Sort all assets by created_at descending
        assets.sort(key=lambda x: x.created_at or "", reverse=True)

        return CardAssetsResponse(
            card_id=card_id, assets=assets, total_count=len(assets)
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error fetching card assets: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card assets retrieval", e),
        ) from e


# ============================================================================
# Velocity
# ============================================================================


@router.get("/cards/{card_id}/velocity")
async def get_card_velocity(
    card_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Get velocity trend summary for a specific card."""
    from app.velocity_service import get_velocity_summary

    summary = await get_velocity_summary(card_id, db)
    if summary is None:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Card not found or velocity data unavailable.",
        )
    return summary


# ============================================================================
# Card Documents
# ============================================================================

# -- Constants ---------------------------------------------------------------

_DOC_MAX_FILE_SIZE = 25 * 1024 * 1024  # 25 MB
_DOC_MAX_EXTRACTED_TEXT = 100_000  # Truncate extracted text to avoid DB bloat

_DOC_ALLOWED_EXTENSIONS: set[str] = {"pdf", "docx", "doc", "txt", "pptx", "xlsx"}

_DOC_ALLOWED_MIME_TYPES: set[str] = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
    "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}

_DOC_VALID_TYPES: set[str] = {
    "nofo",
    "budget",
    "narrative",
    "letter_of_support",
    "application_guide",
    "other",
}


# -- Pydantic response models -----------------------------------------------


class CardDocumentResponse(BaseModel):
    """Representation of a card document (excludes extracted_text for brevity)."""

    id: str
    card_id: str
    uploaded_by: str
    filename: str
    original_filename: str
    blob_path: str
    content_type: str
    file_size_bytes: int
    extraction_status: str
    document_type: str
    description: Optional[str] = None
    metadata: Optional[Dict[str, Any]] = None
    created_at: Optional[str] = None
    updated_at: Optional[str] = None


class CardDocumentListResponse(BaseModel):
    documents: List[CardDocumentResponse]
    total: int


class CardDocumentUploadResponse(BaseModel):
    document: CardDocumentResponse
    message: str


class CardDocumentDownloadUrlResponse(BaseModel):
    url: str
    expires_in_hours: int


# -- Helpers -----------------------------------------------------------------


def _doc_to_response(doc) -> CardDocumentResponse:
    """Convert a CardDocument ORM instance to the API response model."""
    return CardDocumentResponse(
        id=str(doc.id),
        card_id=str(doc.card_id),
        uploaded_by=str(doc.uploaded_by),
        filename=doc.filename,
        original_filename=doc.original_filename,
        blob_path=doc.blob_path,
        content_type=doc.content_type,
        file_size_bytes=doc.file_size_bytes,
        extraction_status=doc.extraction_status,
        document_type=doc.document_type,
        description=doc.description,
        metadata=doc.metadata,
        created_at=doc.created_at.isoformat() if doc.created_at else None,
        updated_at=doc.updated_at.isoformat() if doc.updated_at else None,
    )


def _extract_text(
    data: bytes, content_type: str, original_filename: str
) -> tuple[str, str]:
    """Attempt text extraction from file bytes.

    Returns:
        Tuple of (extracted_text or empty string, extraction_status).
        extraction_status is one of: 'completed', 'failed'.
    """
    ext = (
        original_filename.rsplit(".", 1)[-1].lower() if "." in original_filename else ""
    )

    try:
        if ext == "pdf" or content_type == "application/pdf":
            return _extract_text_from_pdf(data)

        if ext == "txt" or content_type == "text/plain":
            return _extract_text_from_txt(data)

        if ext == "docx" or content_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            return _extract_text_from_docx(data)

        # Unsupported formats (pptx, xlsx, doc): skip extraction
        return "", "failed"

    except Exception as exc:
        logger.warning("Text extraction failed for %s: %s", original_filename, exc)
        return "", "failed"


def _extract_text_from_pdf(data: bytes) -> tuple[str, str]:
    """Extract text from PDF bytes using PyMuPDF (fitz)."""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        logger.warning(
            "PyMuPDF (fitz) not installed -- PDF text extraction unavailable"
        )
        return "", "failed"

    doc = fitz.open(stream=data, filetype="pdf")
    try:
        text_parts: list[str] = []
        max_pages = min(len(doc), 50)
        for page_idx in range(max_pages):
            page = doc[page_idx]
            page_text = page.get_text("text")
            if page_text and page_text.strip():
                text_parts.append(page_text.strip())

        if not text_parts:
            return "", "failed"

        full_text = "\n\n".join(text_parts)
        return full_text[:_DOC_MAX_EXTRACTED_TEXT], "completed"
    finally:
        doc.close()


def _extract_text_from_txt(data: bytes) -> tuple[str, str]:
    """Decode plain text, trying UTF-8 first with Latin-1 fallback."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")

    text = text.strip()
    if not text:
        return "", "failed"
    return text[:_DOC_MAX_EXTRACTED_TEXT], "completed"


def _extract_text_from_docx(data: bytes) -> tuple[str, str]:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.warning("python-docx not installed -- DOCX text extraction unavailable")
        return "", "failed"

    doc = DocxDocument(io.BytesIO(data))
    text_parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    if not text_parts:
        return "", "failed"

    full_text = "\n\n".join(text_parts)
    return full_text[:_DOC_MAX_EXTRACTED_TEXT], "completed"


# -- Endpoints ---------------------------------------------------------------


@router.post(
    "/cards/{card_id}/documents",
    response_model=CardDocumentUploadResponse,
    status_code=status.HTTP_201_CREATED,
)
async def upload_card_document(
    card_id: str,
    file: UploadFile = File(...),
    document_type: str = Form("other"),
    description: Optional[str] = Form(None),
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Upload a document to a card.

    Accepts a multipart form with the file and optional metadata.
    Supported formats: pdf, docx, doc, txt, pptx, xlsx (max 25 MB).
    Text is automatically extracted from PDF, TXT, and DOCX files.

    Args:
        card_id: UUID of the card.
        file: The document file (multipart).
        document_type: One of nofo, budget, narrative, letter_of_support,
                       application_guide, other (default 'other').
        description: Optional human-readable description.

    Returns:
        CardDocumentUploadResponse with the created document metadata.
    """
    from app.models.db.card_document import CardDocument
    from app.storage import attachment_storage

    # Verify card exists
    try:
        card_result = await db.execute(select(Card.id).where(Card.id == card_id))
        if card_result.scalar_one_or_none() is None:
            raise HTTPException(status_code=404, detail="Card not found")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card lookup", e),
        ) from e

    # Validate file
    original_filename = file.filename or "unnamed_file"
    ext = (
        original_filename.rsplit(".", 1)[-1].lower() if "." in original_filename else ""
    )
    if ext not in _DOC_ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"File type '.{ext}' is not allowed. "
                f"Accepted types: {', '.join(sorted(_DOC_ALLOWED_EXTENSIONS))}"
            ),
        )

    content_type = (file.content_type or "application/octet-stream").lower()
    if content_type not in _DOC_ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"MIME type '{content_type}' is not allowed. "
                f"Accepted types: {', '.join(sorted(_DOC_ALLOWED_MIME_TYPES))}"
            ),
        )

    data = await file.read()
    file_size_bytes = len(data)

    if file_size_bytes == 0:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Uploaded file is empty.",
        )
    if file_size_bytes > _DOC_MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=(
                f"File size ({file_size_bytes:,} bytes) exceeds the "
                f"{_DOC_MAX_FILE_SIZE // (1024 * 1024)} MB limit."
            ),
        )

    # Validate document_type
    safe_document_type = document_type if document_type in _DOC_VALID_TYPES else "other"

    # Generate unique stored filename
    unique_prefix = _uuid.uuid4().hex[:8]
    stored_filename = f"{unique_prefix}_{original_filename}"

    try:
        # Upload to blob storage
        blob_path = await attachment_storage.upload_card_document(
            card_id=card_id,
            filename=stored_filename,
            data=data,
            content_type=content_type,
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("document upload", e),
        ) from e

    # Extract text
    extracted_text, extraction_status = _extract_text(
        data, content_type, original_filename
    )

    # Create DB record
    try:
        doc_record = CardDocument(
            card_id=_uuid.UUID(card_id),
            uploaded_by=_uuid.UUID(current_user["id"]),
            filename=stored_filename,
            original_filename=original_filename,
            blob_path=blob_path,
            content_type=content_type,
            file_size_bytes=file_size_bytes,
            extracted_text=extracted_text or None,
            extraction_status=extraction_status,
            document_type=safe_document_type,
            description=description,
        )
        db.add(doc_record)
        await db.flush()
        await db.refresh(doc_record)
    except Exception as e:
        # Best-effort cleanup of uploaded blob on DB failure
        try:
            await attachment_storage.delete(blob_path)
        except Exception:
            logger.warning("Failed to clean up blob %s after DB error", blob_path)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("document record creation", e),
        ) from e

    logger.info(
        "Uploaded card document %s (%s, %d bytes, extraction=%s) for card %s",
        doc_record.id,
        original_filename,
        file_size_bytes,
        extraction_status,
        card_id,
    )

    return CardDocumentUploadResponse(
        document=_doc_to_response(doc_record),
        message=f"Document '{original_filename}' uploaded successfully.",
    )


@router.get(
    "/cards/{card_id}/documents",
    response_model=CardDocumentListResponse,
)
async def list_card_documents(
    card_id: str,
    document_type: Optional[str] = Query(None, description="Filter by document type"),
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """List all documents for a card.

    Returns document metadata without the full extracted_text field
    to keep responses compact.

    Args:
        card_id: UUID of the card.
        document_type: Optional filter by document type.

    Returns:
        CardDocumentListResponse with list of documents and total count.
    """
    from app.models.db.card_document import CardDocument

    try:
        # Verify card exists
        card_result = await db.execute(select(Card.id).where(Card.id == card_id))
        if card_result.scalar_one_or_none() is None:
            raise HTTPException(status_code=404, detail="Card not found")

        stmt = (
            select(CardDocument)
            .where(CardDocument.card_id == card_id)
            .order_by(CardDocument.created_at.desc())
        )

        if document_type and document_type in _DOC_VALID_TYPES:
            stmt = stmt.where(CardDocument.document_type == document_type)

        result = await db.execute(stmt)
        docs = result.scalars().all()

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("card documents listing", e),
        ) from e

    responses = [_doc_to_response(d) for d in docs]
    return CardDocumentListResponse(documents=responses, total=len(responses))


@router.delete(
    "/cards/{card_id}/documents/{document_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
async def delete_card_document(
    card_id: str,
    document_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Delete a document from a card.

    Removes the file from blob storage and the database record.

    Args:
        card_id: UUID of the card.
        document_id: UUID of the document.

    Raises:
        HTTPException 404: Document not found or does not belong to this card.
    """
    from app.models.db.card_document import CardDocument
    from app.storage import attachment_storage

    try:
        result = await db.execute(
            select(CardDocument).where(
                CardDocument.id == document_id,
                CardDocument.card_id == card_id,
            )
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found",
            )

        # Only the uploader may delete a document
        if str(doc.uploaded_by) != current_user["id"]:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="You can only delete documents you uploaded",
            )

        blob_path = doc.blob_path

        # Delete DB record first
        await db.delete(doc)
        await db.flush()

        # Delete from blob storage
        try:
            await attachment_storage.delete(blob_path)
        except Exception as blob_err:
            logger.warning(
                "Failed to delete blob %s (DB record already removed): %s",
                blob_path,
                blob_err,
            )

        logger.info(
            "Deleted card document %s (blob: %s) by user %s",
            document_id,
            blob_path,
            current_user["id"],
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("document deletion", e),
        ) from e


@router.get(
    "/cards/{card_id}/documents/{document_id}/download",
    response_model=CardDocumentDownloadUrlResponse,
)
async def get_card_document_download_url(
    card_id: str,
    document_id: str,
    current_user: dict = Depends(get_current_user_hardcoded),
    db: AsyncSession = Depends(get_db),
):
    """Get a time-limited download URL for a card document.

    Generates a SAS URL valid for 1 hour.

    Args:
        card_id: UUID of the card.
        document_id: UUID of the document.

    Returns:
        CardDocumentDownloadUrlResponse with the SAS URL and expiry info.

    Raises:
        HTTPException 404: Document not found or does not belong to this card.
    """
    from app.models.db.card_document import CardDocument
    from app.storage import attachment_storage

    try:
        result = await db.execute(
            select(CardDocument).where(
                CardDocument.id == document_id,
                CardDocument.card_id == card_id,
            )
        )
        doc = result.scalar_one_or_none()
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found",
            )

        url = await attachment_storage.generate_sas_url(doc.blob_path)
        return CardDocumentDownloadUrlResponse(url=url, expires_in_hours=1)

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=_safe_error("document download URL generation", e),
        ) from e
